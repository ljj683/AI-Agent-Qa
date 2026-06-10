"""上传文档问答引擎

用于 Streamlit 上传任意 PDF / DOCX / TXT / MD 后即时问答。
实现流程：文档解析 -> 文本清洗 -> 分块 -> TF-IDF 向量化 -> 余弦相似度检索 -> DeepSeek 回答生成。
"""
from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Dict, Any

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class FileSummary:
    file_name: str
    file_type: str
    file_size_mb: float
    text_length: int
    chunk_count: int


def clean_text(text: str) -> str:
    """基础文本清洗：去控制字符、合并多余空白，保留可读内容。"""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def split_text(text: str, chunk_size: int = 800, overlap: int = 120) -> List[str]:
    """滑动窗口分块，适合中文和英文混合文档。"""
    text = clean_text(text)
    if not text:
        return []
    if chunk_size <= overlap:
        raise ValueError("chunk_size 必须大于 overlap")

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if len(chunk) >= 30:
            chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def _decode_text(data: bytes) -> str:
    for enc in ["utf-8", "utf-8-sig", "gb18030", "gbk", "latin1"]:
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_text_from_bytes(file_name: str, data: bytes) -> str:
    """根据后缀解析上传文件。"""
    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("暂不支持该文件类型，请上传 PDF、DOCX、TXT 或 MD 文件。")

    if suffix == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=data, filetype="pdf")
            pages = []
            for page in doc:
                pages.append(page.get_text("text"))
            text = "\n".join(pages)
        except Exception as exc:
            raise RuntimeError(f"PDF 解析失败：{exc}") from exc

    elif suffix == ".docx":
        try:
            from docx import Document
            document = Document(io.BytesIO(data))
            paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
            # 兼容读取简单表格文本
            table_texts = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_texts.append(" | ".join(cells))
            text = "\n".join(paragraphs + table_texts)
        except Exception as exc:
            raise RuntimeError(f"Word 解析失败：{exc}") from exc

    else:  # .txt / .md
        text = _decode_text(data)

    text = clean_text(text)
    if len(text) < 30:
        raise RuntimeError(
            "文档解析出的文本过少。可能是扫描版 PDF、图片型 Word，或文件内容为空。"
            "扫描版 PDF 建议先用 MinerU/OCR 转成 Markdown 或 TXT 后再上传。"
        )
    return text


class UploadedDocumentIndex:
    """临时上传文档检索索引。支持单文档或多文档同时上传。"""

    def __init__(self, chunks: List[Dict[str, Any]], summaries: List[FileSummary]):
        if not chunks:
            raise ValueError("没有可检索的文本块。")
        self.chunks = chunks
        self.summaries = summaries
        # char ngram 对中文更友好，不依赖分词器，适合轻量部署。
        self.vectorizer = TfidfVectorizer(analyzer="char",ngram_range=(2, 4),min_df=1,max_df=0.95)
        self.matrix = self.vectorizer.fit_transform([c["text"] for c in chunks])

    @classmethod
    def from_uploaded_files(cls, uploaded_files: Iterable[Any], chunk_size: int = 800, overlap: int = 120):
        all_chunks: List[Dict[str, Any]] = []
        summaries: List[FileSummary] = []

        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            file_data = uploaded_file.getvalue()
            suffix = Path(file_name).suffix.lower()
            text = extract_text_from_bytes(file_name, file_data)
            chunks = split_text(text, chunk_size=chunk_size, overlap=overlap)

            summaries.append(
                FileSummary(
                    file_name=file_name,
                    file_type=suffix.replace(".", "").upper(),
                    file_size_mb=round(len(file_data) / 1024 / 1024, 2),
                    text_length=len(text),
                    chunk_count=len(chunks),
                )
            )

            for i, chunk in enumerate(chunks):
                all_chunks.append({
                    "chunk_id": f"{Path(file_name).stem}_{i}",
                    "source_file": file_name,
                    "text": chunk,
                    "length": len(chunk),
                })

        return cls(all_chunks, summaries)

    def retrieve(self, question: str, top_k: int = 5) -> List[Dict[str, Any]]:
        query_vec = self.vectorizer.transform([question])
        scores = cosine_similarity(query_vec, self.matrix).ravel()
        top_indices = scores.argsort()[::-1][:top_k]
        results = []
        for idx in top_indices:
            item = dict(self.chunks[int(idx)])
            item["score"] = float(scores[int(idx)])
            results.append(item)
        return results


def generate_answer_with_deepseek(question: str, contexts: List[Dict[str, Any]], history: list | None = None) -> str:
    """根据上传文档检索片段调用 DeepSeek 生成回答。无 Key 时返回本地降级答案。"""
    api_key = os.getenv("DEEPSEEK_API_KEY", "")
    base_url = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
    model = os.getenv("DEEPSEEK_MODEL", "deepseek-v4-flash")

    context_text = "\n\n".join(
        f"片段{i + 1}｜来源：{ctx.get('source_file', '上传文档')}｜相似度：{ctx.get('score', 0):.4f}\n{ctx.get('text', '')}"
        for i, ctx in enumerate(contexts)
    )

    history_text = ""
    if history:
        for item in history[-3:]:
            history_text += f"用户：{item.get('question', '')}\n助手：{item.get('answer', '')}\n"

    if not contexts or max([c.get("score", 0) for c in contexts]) <= 0:
        return "上传文档中没有找到明确依据。请换一个更贴近文档内容的问题，或检查文档是否解析成功。"

    if not api_key:
        return (
            "当前未配置 DeepSeek API Key，以下为基于上传文档检索片段的本地降级回答：\n\n"
            + "\n\n".join(
                f"【{i + 1}】来源：{ctx.get('source_file', '上传文档')}｜相似度：{ctx.get('score', 0):.4f}\n{ctx.get('text', '')[:600]}"
                for i, ctx in enumerate(contexts)
            )
        )

    prompt = f"""
你是一个专业、严谨的上传文档问答 AI 智能体。请严格根据“上传文档检索片段”回答用户问题。

回答要求：
1. 先给出简洁结论，再分点解释；
2. 必须优先依据上传文档内容回答；
3. 如果上传文档中没有明确依据，请说明“上传文档中没有找到明确依据”，不要编造；
4. 回答要适合学生理解，结构清晰；
5. 回答结尾写“参考依据：上传文档检索片段”。

历史对话：
{history_text}

用户问题：
{question}

上传文档检索片段：
{context_text}
"""

    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url=base_url)
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你是基于 RAG 的文档问答助手，回答必须基于检索内容。"},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
        )
        return response.choices[0].message.content
    except Exception as exc:
        return (
            f"DeepSeek 调用失败，系统已切换为本地检索降级回答。\n\n错误原因：{exc}\n\n"
            + "\n\n".join(
                f"【{i + 1}】来源：{ctx.get('source_file', '上传文档')}｜相似度：{ctx.get('score', 0):.4f}\n{ctx.get('text', '')[:600]}"
                for i, ctx in enumerate(contexts)
            )
        )
